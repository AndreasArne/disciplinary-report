import json
import sys
import pyperclip
from docx import Document

def get_unique_cells_ordered(table):
    """
    Rows with less cells than max on other rows
    get duplicated cells to match number of cells as max.
    Filter the duplicate cells and only get unique.
    """
    unique = []
    for row in table.rows:
        for cell in row.cells:
            if cell not in unique:
                unique.append(cell)
    return unique



def insert_student_data(table):
    """
    Take student data from clipboard and add to table
    """
    student = pyperclip.paste().split("\n")
    cells = get_unique_cells_ordered(table)

    for cell, stud_row in zip(cells, student):
        cell.add_paragraph(stud_row.strip())



def order_data(date, code, course, teacher):
    """
    Put data in order of cell occurance
    """
    return [
        date,
        course["name"],
        code,
        course["part_name"],
        course["points"],
        teacher["name"],
        teacher["phone"],
        teacher["section"]
    ]



def insert_offense_data(table, data):
    """
    Add offence data
    """
    cells = get_unique_cells_ordered(table)
    for cell, stud_row in zip(cells, data):
        cell.add_paragraph(stud_row.strip())



def insert_appendix_data(table, args):
    cells = table.column_cells(0)
    stud_ladok = args[3] + "-ladok.pdf"
    courseplan = f"kursplan-{args[1]}.pdf"
    description = "Beskrivning av ärende.pdf"
    instructions = args[4]
    extras = args[5]

    cells[0].add_paragraph(stud_ladok)
    cells[1].add_paragraph(courseplan)
    cells[2].add_paragraph(description)
    cells[4].add_paragraph(instructions)
    if len(args) == 6:
        cells[5].add_paragraph(extras)

    expected_files = (
        "Expects you to have/create the following files:",
        stud_ladok,
        courseplan,
        description,
        instructions,
        extras.replace(",", "\n\t")
    )
    return expected_files


def get_argvs():
    if len(sys.argv) not in (6, 7): # +1 for filename
        print(
            "Need the following arguments in order:",
            "   Date of offence",
            "   Course code",
            "   Teacher acronym",
            "   Student acronym",
            "   Assignment instructions filename/url",
            "   optionally add a comma seperated list (no spaces) of additional files used in report",
            "",
            "   Exampe: python3 main.py 2020-11-17 DV1531 aar anar12 https://dbwebb.se code.zip,video.mp4",
            sep="\n"
        )
        sys.exit(1)
    return sys.argv[1:]



def main():
    args = get_argvs()
    doc = Document("template/report.docx")
    insert_student_data(doc.tables[0])

    admin_data = json.load(open("local-data.json"))
    ordered_data = order_data(
        args[0],
        args[1],
        admin_data["courses"][args[1]],
        admin_data["teachers"][args[2]]
    )
    insert_offense_data(doc.tables[1], ordered_data)

    expected_files = insert_appendix_data(doc.tables[2], args)

    output = f"output/Anmälan om disciplinförseelse {args[3]}.docx"
    doc.save(output)
    
    print(
        "",
        f"Convert '{output}' to pdf and then sign it online.",
        "",
        "\n\t".join(expected_files),
        "",
        "Förslag på text att skicka till studenten. GLÖM INTE ATT ERSÄTTA <uppgifter>",
        '"""',
        admin_data["text"].format(teacher=admin_data["teachers"][args[2]]["name"]),
        '"""',
        f"Link for downloading courseplan https://edu.bth.se/utbildning/utb_kursplaner.asp?KKurskod={args[1]}"
        "",
        sep="\n"
    )

if __name__ == "__main__":
    main()
